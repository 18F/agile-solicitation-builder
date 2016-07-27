# -*- coding: utf-8 -*-
import datetime

from docx import Document
from api.models import Agency, RFQ, ContentComponent, Deliverable, CustomComponent, session


BIG_HEADING = 1
SUB_HEADING = 2
SMALL_HEADING = 4

user_dict = {
    "external_people": "External People/The Public",
    "external_it": "External IT/Developers",
    "internal_people": "Internal People/Government Employees",
    "internal_it": "Internal IT/Developers",
}


def make_dict(components):
    component_dict = {}
    for component in components:
        component_dict[component.name] = component.text
    return component_dict


def make_custom_component_list(components):
    custom_component_list = []
    for component in components:
        this_component = {}
        this_component['name'] = component.name
        this_component['text'] = component.text
        this_component['title'] = component.title
        custom_component_list.append(this_component)

    return custom_component_list


def get_users(cc, user_types):
    users = []
    for user in user_types:
        if cc[user] == "true":
            users.append(user)
    return users

# Add global counter for headings
# def section_heading(document, headingTitle):
#     title = str(sectionNumber)+'. '+headingTitle
#     document.add_heading(headingTitle, level=BIG_HEADING)
#     global sectionNumberCounter
#     sectionNumberCounter = sectionNumberCounter+1
#     return document

def overview(document, rfq):
    # table of contents & basic info

    agency_full_name = session.query(Agency).filter_by(abbreviation=rfq.agency).first().full_name
    title = "RFQ for the " + agency_full_name
    document.add_heading(title, level=BIG_HEADING)
    doc_date = str(datetime.date.today())
    document.add_heading(doc_date, level=3)

    # table of contents
    document.add_heading("Table of Contents", level=SUB_HEADING)
    sections = ["Definitions", "Services", "Statement of Objectives", "Personnel Requirements", "Inspection and Delivery", "Government Roles", "Special Requirements", "Additional Contract Clauses", "Appendix"]
    for section in sections:
        document.add_paragraph(section, style='ListNumber')

    text = "Note: All sections of this RFQ will be incorporated into the contract except the Statement of Objectives, Instructions, and Evaluation Factors."
    document.add_paragraph(text)
    document.add_page_break()

    return document


def definitions(document, rfq):

    document.add_heading("1. Definitions", level=BIG_HEADING)
    all_definitions = session.query(ContentComponent).filter_by(document_id=rfq.id).filter_by(section=1).first()
    for definition in all_definitions.text.split("\n\n"):
        document.add_paragraph(definition)

    return document


def services(document, rfq):
    document.add_heading("2. Services", level=BIG_HEADING)
    content_components = session.query(ContentComponent).filter_by(document_id=rfq.id).filter_by(section=2).all()
    # include vendor number
    cc = make_dict(content_components)
    optionPeriods = cc["optionPeriods"]
    document.add_heading("Brief Description of Services & Type of Contract", level=SUB_HEADING)
    document.add_paragraph(cc["descriptionOfServices"])
    document.add_paragraph(cc["naicsText"])

    document.add_heading("Budget", level=SUB_HEADING)
    max_text = "The government is willing to invest a maximum budget of $" + cc["maxBudget"] + " in this endeavor."
    document.add_paragraph(max_text)

    # travel
    if cc["travelRequirement"] == "yes":
        travel_text = "The Government anticipates travel will be required under this effort. Contractor travel expenses will not exceed $" + cc["travelBudget"] + "."
        document.add_paragraph(travel_text)
        document.add_paragraph(cc["travelLanguage"])
    else:
        document.add_paragraph("The Government does not anticipate significant travel under this effort.")

    # @TODO make top column bold, add award fee/incentive information (if applicable)
    # base period
    document.add_heading("Contract Line Item Number (CLIN) Format", level=SUB_HEADING)
    document.add_paragraph("\n")

    table = document.add_table(rows=2, cols=1)
    table.style = 'TableGrid'
    table.rows[0].cells[0].text = "Base Period: " + str(cc["basePeriodDurationNumber"]) + ' ' + cc["basePeriodDurationUnit"]
    table.rows[1].cells[0].text = "CLIN 0001, FFP- Completion - The Contractor shall provide services for the Government in accordance with the Performance Work Statement (PWS)"

    table = document.add_table(rows=4, cols=2)
    table.style = 'TableGrid'
    table.rows[0].cells[0].text = "Iteration Period of Performance"
    table.rows[0].cells[1].text = cc["iterationPoPNumber"] + ' ' + cc["iterationPoPUnit"]
    table.rows[1].cells[0].text = "Price Per Iteration"
    table.rows[1].cells[1].text = "$XXXXX (Vendor Completes)"
    table.rows[2].cells[0].text = "Period of Performance"
    table.rows[2].cells[1].text = cc["basePeriodDurationNumber"] + ' ' + cc["basePeriodDurationUnit"]
    table.rows[3].cells[0].text = "Firm Fixed Price (Completion):"
    table.rows[3].cells[1].text = "$XXXXX (Vendor Completes)"
    # @TODO if base fee, add base fee clin row

    document.add_paragraph("\n")

    # option periods
    for i in range(1, int(optionPeriods)+1):
        table = document.add_table(rows=2, cols=1)
        table.style = 'TableGrid'
        table.rows[0].cells[0].text = "Option Period " + str(i) + ": " + str(cc["optionPeriodDurationNumber"]) + ' ' + cc["optionPeriodDurationUnit"]
        table.rows[1].cells[0].text = "CLIN " + str(i) + "0001, FFP- Completion - The Contractor shall provide services for the Government in accordance with the Performance Work Statement (PWS)"

        table = document.add_table(rows=4, cols=2)
        table.style = 'TableGrid'
        table.rows[0].cells[0].text = "Iteration Period of Performance"
        table.rows[0].cells[1].text = cc["iterationPoPNumber"] + ' ' + cc["iterationPoPUnit"]
        table.rows[1].cells[0].text = "Price Per Iteration"
        table.rows[1].cells[1].text = "$XXXXX (Vendor Completes)"
        table.rows[2].cells[0].text = "Period of Performance"
        table.rows[2].cells[1].text = cc["optionPeriodDurationNumber"] + ' ' + cc["optionPeriodDurationUnit"]
        table.rows[3].cells[0].text = "Firm Fixed Price (Completion):"
        table.rows[3].cells[1].text = "$XXXXX (Vendor Completes)"
        # @TODO if option fee, add option fee clin row

        document.add_paragraph("\n")

    # @TODO add custom CLIN(s)

    document.add_heading("Payment Schedule", level=SUB_HEADING)
    document.add_paragraph(cc["paymentSchedule"])

    return document


def objectives(document, rfq):
    document.add_heading("3. Objectives", level=BIG_HEADING)
    content_components = session.query(ContentComponent).filter_by(document_id=rfq.id).filter_by(section=3).all()
    cc = make_dict(content_components)
    document.add_heading("General Background", level=SUB_HEADING)
    if len(cc["generalBackground"]) > 0:
        document.add_paragraph(cc["generalBackground"])
    else:
        document.add_paragraph("Please provide several paragraphs about your project's history, mission, and current state.")

    document.add_heading("Program History", level=SUB_HEADING)
    if len(cc["programHistory"]) > 0:
        document.add_paragraph(cc["programHistory"])
    else:
        document.add_paragraph("If you have any information about the current vendors and specific technology being used please provide it here.")

    document.add_heading("Users", level=SUB_HEADING)
    user_types = ["external_people", "external_it", "internal_people", "internal_it"]
    users = get_users(cc, user_types)
    if len(users) == 0:
        document.add_paragraph("The primary users may include any of the following:")
        for i, user in enumerate(user_dict):
            document.add_paragraph(str(i+1) + ".  " + user_dict[user])

    else:
        document.add_paragraph("The primary users will include the following:")
        for i, user in enumerate(users):
            document.add_paragraph(str(i+1) + ".  " + user_dict[user])

    # for user in users, add text of each user's needs
    needs = ['external_people_needs', 'external_it_needs', 'internal_it_needs', 'internal_people_needs']

    document.add_heading("User Research", level=SUB_HEADING)
    user_research_options = {
        "done": "Research has already been conducted, either internally or by another vendor.",
        "internal": "We intend to conduct user research internally prior to the start date of this engagement.",
        "vendor": "The vendor will be responsible for the user research.",
    }

    if cc["userResearchStrategy"] == "vendor":
        document.add_paragraph(user_research_options["vendor"])
        document.add_heading("Understand What People Need", level=SUB_HEADING)
        document.add_paragraph(cc["whatPeopleNeed"])

        document.add_heading("Address the whole experience, from start to finish", level=SUB_HEADING)
        document.add_paragraph(cc["startToFinish"])

    if cc["userResearchStrategy"] == "done":
        document.add_paragraph(user_research_options["done"])

    if cc["userResearchStrategy"] == "internal":
        document.add_paragraph(user_research_options["internal"])

    if cc["userResearchStrategy"] == "none":
        pass

    document.add_paragraph(cc['userAccess'])

    document.add_heading("Universal Requirements", level=BIG_HEADING)

    document.add_heading("Build the service using agile and iterative practices", level=SUB_HEADING)
    document.add_paragraph(cc["agileIterativePractices"])

    document.add_heading("Make it simple and intuitive", level=SUB_HEADING)
    document.add_paragraph(cc["simpleAndIntuitive"])

    document.add_heading("Use data to drive decisions", level=SUB_HEADING)
    document.add_paragraph(cc["dataDrivenDecisions"])

    document.add_heading("Specific Tasks and Deliverables", level=SUB_HEADING)
    text = "This " + rfq.doc_type + " will require the following services:"
    document.add_paragraph(text)

    deliverables = session.query(Deliverable).filter_by(document_id=rfq.id).filter_by(value="true").all()
    for deliverable in deliverables:
        document.add_paragraph("    " + deliverable.display)


    document.add_heading("Deliverables", level=SUB_HEADING)
    document.add_paragraph(cc["definitionOfDone"])

    for deliverable in deliverables:
        document.add_heading(deliverable.display, level=SMALL_HEADING)
        document.add_paragraph(deliverable.text)

    document.add_heading("Place of Performance", level=SUB_HEADING)
    if cc['locationRequirement'] == "no":
        document.add_paragraph("The contractor is not required to have a full-time working staff presence on-site.")

    else:
        if len(cc['locationText']) > 0:
            location = cc["locationText"]
        else:
            location = "[LOCATION HERE]"

        location_text = "The contractor shall have a full-time working staff presence at " + location + ". The contractor shall have additional facilities to perform contract functions as necessary."
        document.add_paragraph(location_text)
        document.add_paragraph(cc["offSiteDevelopmentCompliance"])

    document.add_heading("Kick Off Meeting", level=SUB_HEADING)
    kickoff_text = ""
    if cc["kickOffMeeting"] == "none":
        kickoff_text = "A formal kick-off meeting will not be required."
    if cc["kickOffMeeting"] == "in-person":
        kickoff_text = cc["kickOffMeetingInPerson"]
    if cc["kickOffMeeting"] == "remote":
        kickoff_text = cc["kickOffMeetingRemote"]

    document.add_paragraph(kickoff_text)

    return document


def personnel(document, rfq):
    document.add_heading("4. Key Personnel", level=BIG_HEADING)
    content_components = session.query(ContentComponent).filter_by(document_id=rfq.id).filter_by(section=4).all()
    cc = make_dict(content_components)

    document.add_paragraph(cc["keyPersonnelIntro"])

    document.add_heading("Security Clearance and Onsite Presence", level=SUB_HEADING)
    if cc["clearanceRequired"] == "None":
        document.add_paragraph("Contractor personnel will not be required to have a security clearance.")
    else:
        document.add_paragraph("Some contractor personnel will be required to have a clearance at the level of " + cc["clearanceRequired"] + ".")

    if cc["onSiteRequired"] == "yes":
        document.add_paragraph("An onsite presence by the contractor will be required.")
    else:
        document.add_paragraph("An onsite presence by the contractor will not be required.")

    document.add_heading("Key Personnel Evaluation", level=SUB_HEADING)

    if cc["evaluateKeyPersonnel"] == "yes":
        document.add_paragraph(cc["keyPersonnelRequirements"])
    else:
        document.add_paragraph(cc["notEvaluateKeyPersonnel"])

    document.add_heading("Performance Work Statement", level=SUB_HEADING)
    document.add_paragraph(cc["performanceWorkStatement"])

    return document


def invoicing(document, rfq):
    document.add_heading("5. Invoicing & Funding", level=BIG_HEADING)

    content_components = session.query(ContentComponent).filter_by(document_id=rfq.id).filter_by(section=5).all()
    cc = make_dict(content_components)

    document.add_paragraph(cc["invoicing"])

    document.add_paragraph("The Contractor shall submit an original invoice for payment to the following office:")
    if len(cc['billingAddress']) < 1:
        document.add_paragraph("ADD BILLING ADDRESS HERE")
    else:
        document.add_paragraph(cc["billingAddress"])
    document.add_paragraph(cc["duplicateInvoice"])

    return document


def inspection_and_delivery(document, rfq):
    document.add_heading("6. Inspection and Delivery", level=BIG_HEADING)

    content_components = session.query(ContentComponent).filter_by(document_id=rfq.id).filter_by(section=6).all()
    cc = make_dict(content_components)

    document.add_heading("Overview", level=SUB_HEADING)
    document.add_paragraph(cc["guidingPrinciples"])

    document.add_heading("Delivery and Timing", level=SUB_HEADING)
    document.add_paragraph(cc["inspectionOverview"])

    document.add_heading("Late Delivery", level=SUB_HEADING)
    document.add_paragraph(cc["lateDelivery"])

    document.add_heading("Collaboration Environment", level=SUB_HEADING)

    if cc["workspaceExists"] == "yes":
        if len(cc["workspaceName"]) > 0:
            document.add_paragraph(rfq.agency + " is currently using " + cc["workspaceName"] + " as their primary collaborative workspace tool. The contractor is required to establish a collaborative workspace using either this tool or another that both the contractor and the CO can agree upon.")

    document.add_paragraph(cc["deliveringDeliverables"])

    document.add_heading("Transition Activities", level=SUB_HEADING)
    document.add_paragraph(cc["transitionActivities"])

    return document


def government_roles(document, rfq):
    document.add_heading("7. Government Roles", level=BIG_HEADING)

    content_components = session.query(ContentComponent).filter_by(document_id=rfq.id).filter_by(section=7).all()
    cc = make_dict(content_components)

    custom_components = session.query(CustomComponent).filter_by(document_id=rfq.id).filter_by(section=7).order_by(CustomComponent.id).all()

    document.add_paragraph(cc["stakeholderIntro"])

    component_list = make_custom_component_list(custom_components)

    for component in component_list:
        document.add_heading(component['title'], level=SUB_HEADING)
        document.add_paragraph(component['text'])

    return document


def special_requirements(document, rfq):
    document.add_heading("8. Special Requirements", level=BIG_HEADING)

    custom_components = session.query(CustomComponent).filter_by(document_id=rfq.id).filter_by(section=8).order_by(CustomComponent.id).all()

    component_list = make_custom_component_list(custom_components)

    for component in component_list:
        document.add_heading(component['title'], level=SUB_HEADING)
        document.add_paragraph(component['text'])

    return document


def contract_clauses(document, rfq):
    document.add_heading("9. Additional Contract Clauses", level=BIG_HEADING)
    contract_clauses = session.query(ContentComponent).filter_by(document_id=rfq.id).filter_by(section=9).first()
    document.add_paragraph(contract_clauses.text)

    return document

def instructions_to_offerors(document, rfq):
    document.add_heading("10. Instructions to Offerors", level=BIG_HEADING)
    instructions = session.query(ContentComponent).filter_by(document_id=rfq.id).filter_by(section=10).first()
    document.add_paragraph(instructions.text)

    return document

def evaluation_criteria(document, rfq):
    document.add_heading("11. Evaluation Criteria", level=BIG_HEADING)
    instructions = session.query(ContentComponent).filter_by(document_id=rfq.id).filter_by(section=11).first()
    document.add_paragraph(instructions.text)

    return document

def appendix(document, rfq):

    document.add_heading("12. Appendix", level=BIG_HEADING)

    return document


def create_document(rfq_id):
    rfq = session.query(RFQ).filter_by(id=rfq_id).first()

    document = Document()

    document = overview(document, rfq)
    document = definitions(document, rfq)
    document = services(document, rfq)
    document = objectives(document, rfq)
    document = personnel(document, rfq)
    document = invoicing(document, rfq)
    document = inspection_and_delivery(document, rfq)
    document = government_roles(document, rfq)
    document = special_requirements(document, rfq)
    document = contract_clauses(document, rfq)
    document = instructions_to_offerors(document, rfq)
    document = evaluation_criteria(document, rfq)
    document = appendix(document, rfq)

    return document
